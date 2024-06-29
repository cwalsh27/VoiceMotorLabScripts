import os
from docx import Document

desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
doc_path = os.path.join(desktop_path, r'TextFiles/CSSPRING20230512-181729_Recording_1280x720_MM.docx')
new_doc_path = os.path.join(desktop_path, r'TextFiles/outputFile')

doc = Document(doc_path)
newDoc = Document()

# appending line labels
last_speaker_B = False
lines = []
count = 0

# for para in doc.paragraphs:
#     for run in para.runs:
#         print(run.text, "\n")

# original draft

skipChars = [".", ",", " ", "[", "]"]


def check_valid_run(run):
    valid_char = False
    for char in run.text:
        if char not in skipChars:
            valid_char = True
    if valid_char:
        return True
    else:
        return False


for para in doc.paragraphs:
    newP = newDoc.add_paragraph('')
    if para.text[0:2] == "((" and para.text[-2:] == "))":
        newP.add_run(para.text)
    else:
        if any((run.bold and check_valid_run(run)) for run in para.runs) and not last_speaker_B:
            newP.add_run(f'B{count + 1}: {para.text}').bold = True
            last_speaker_B = True
            count += 1
        elif any((run.bold and check_valid_run(run)) for run in para.runs) and last_speaker_B:
            newP.add_run(para.text).bold = True
            # count += 1
        elif last_speaker_B:
            newP.add_run(f'A{count + 1}: {para.text}')
            last_speaker_B = False
            count += 1
        else:
            newP.add_run(para.text)
            # count += 1

newDoc.save('newfile.docx')
